from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, validator
import subprocess
import os
import json
import re
import uuid
import asyncio
from datetime import datetime

app = FastAPI(title="PhoenixForge 🔥")

# Global dict to store background task progress
active_tasks = {}

class IdeaRequest(BaseModel):
    idea: str = Field(..., min_length=3, max_length=150)
    
    @validator('idea')
    def sanitize_idea(cls, v):
        # Reject HTML tags or script symbols to prevent XSS/HTML injection
        if '<' in v or '>' in v:
            raise ValueError('Idea cannot contain HTML/Script tag characters (< or >).')
        # Allow only alphanumeric, spaces, and basic punctuation
        sanitized = re.sub(r'[^a-zA-Z0-9\s\-_,\.!\?\'"()]', '', v)
        if len(sanitized.strip()) < 3:
            raise ValueError('Idea must contain at least 3 safe characters.')
        return sanitized.strip()

def extract_pipeline_metadata(content):
    pipeline = {"gathered": "N/A", "cleaned": "N/A", "organized": "N/A", "presented": "N/A"}
    pipeline_match = re.search(r'## Pipeline Metadata\s*```json\s*([\s\S]*?)```', content)
    if not pipeline_match:
        pipeline_match = re.search(r'## Pipeline Metadata\s*([\s\S]*?)(?=\n##|$)', content)
    if pipeline_match:
        try:
            pipeline = json.loads(pipeline_match.group(1).strip())
        except:
            pass
    return pipeline

async def run_pipeline_worker(task_id: str, idea: str):
    import engine
    import coordinator
    try:
        # Consult coordinator to get top performing strategies
        best_strats = coordinator.get_best_strategies(limit=3)
        
        # Generate optimized queries based on past performance
        queries, templates = coordinator.generate_optimized_queries(idea, best_strats)
        
        strategy_metadata = {
            "query_templates": templates,
            "queries_used": queries,
            "scraper_type": "crawl4ai_BS4_fallback",
            "prompt_version": 1
        }
        
        # Run the CPU-heavy scraping & LLM pipeline in a separate thread pool
        # to keep the FastAPI main event loop free for status polling.
        payload = await asyncio.to_thread(
            engine.run_phoenixforge,
            idea,
            task_id,
            active_tasks,
            queries,
            strategy_metadata
        )
        active_tasks[task_id]["status"] = "completed"
        active_tasks[task_id]["current_step"] = "Completed"
        active_tasks[task_id]["message"] = "Analysis report generated successfully."
        active_tasks[task_id]["results"] = payload
    except Exception as e:
        import traceback
        traceback.print_exc()
        active_tasks[task_id]["status"] = "failed"
        active_tasks[task_id]["current_step"] = "Analysis failed"
        active_tasks[task_id]["message"] = str(e)

@app.post("/api/analyze")
async def analyze_risk(request: IdeaRequest):
    if not request.idea:
        raise HTTPException(status_code=400, detail="Idea cannot be empty")
    
    task_id = str(uuid.uuid4())
    active_tasks[task_id] = {
        "status": "processing",
        "current_step": "Initializing search queries...",
        "extracted_count": 0,
        "total_articles": 0,
        "message": "Starting research agent...",
        "results": None
    }
    
    # Schedule task concurrently on the event loop
    asyncio.create_task(run_pipeline_worker(task_id, request.idea))
    
    return {"task_id": task_id}

@app.get("/api/status/{task_id}")
async def get_task_status(task_id: str):
    if task_id not in active_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    return active_tasks[task_id]

@app.get("/api/coordinator/status")
async def get_coordinator_status():
    try:
        import coordinator
        summary = coordinator.summarize_learning()
        best = coordinator.get_best_strategies(limit=1)
        best_score = best[0].get("quality_score", 0.0) if best else 0.0
        return {
            "status": "success",
            "summary": summary,
            "best_score": best_score
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/history/clear")
async def clear_history():
    try:
        import history
        history.clear_all_history()
        return {"status": "success", "message": "History cleared successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def render_mermaid_to_image(mermaid_code):
    import base64
    import urllib.parse
    import requests
    import tempfile
    import os
    
    encoded = urllib.parse.quote(mermaid_code)
    
    # Strategy 1: Try mermaid.ink PNG (online, highly compatible)
    try:
        url = f"https://mermaid.ink/img/{encoded}"
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            return base64.b64encode(response.content).decode('utf-8'), "png"
    except Exception as e:
        print(f"mermaid.ink img failed: {e}")
        
    # Strategy 1b: Try mermaid.ink SVG
    try:
        url = f"https://mermaid.ink/svg/{encoded}"
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            return base64.b64encode(response.content).decode('utf-8'), "svg"
    except Exception as e:
        print(f"mermaid.ink svg failed: {e}")
        
    # Strategy 2: Try local Mermaid CLI (mmdc) to render to SVG
    try:
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', encoding='utf-8', delete=False) as f:
            f.write(mermaid_code)
            input_path = f.name
        output_path = input_path + ".svg"
        
        # Try global mmdc first
        try:
            subprocess.run(
                ['mmdc', '-i', input_path, '-o', output_path, '-b', 'transparent'],
                check=True,
                timeout=10,
                capture_output=True,
                shell=True
            )
        except Exception as cli_err:
            print(f"Direct mmdc failed, trying npx: {cli_err}")
            # Try via npx
            subprocess.run(
                ['npx', '--yes', '@mermaid-js/mermaid-cli', 'mmdc', '-i', input_path, '-o', output_path, '-b', 'transparent'],
                check=True,
                timeout=15,
                capture_output=True,
                shell=True
            )
            
        with open(output_path, 'rb') as f:
            svg_data = f.read()
            
        os.unlink(input_path)
        os.unlink(output_path)
        return base64.b64encode(svg_data).decode('utf-8'), "svg"
    except Exception as e:
        print(f"Local Mermaid CLI failed: {e}")
        
    # Strategy 3: Fallback placeholder (a clean text box SVG)
    placeholder_svg = f'''
    <svg width="400" height="100" xmlns="http://www.w3.org/2000/svg">
      <rect width="100%" height="100%" fill="#f8f9fa" rx="8" />
      <text x="50%" y="45%" font-family="Arial" font-size="16" fill="#6c757d" text-anchor="middle">
        📊 Flowchart rendering unavailable
      </text>
      <text x="50%" y="65%" font-family="Arial" font-size="12" fill="#adb5bd" text-anchor="middle">
        Please view the live dashboard at http://127.0.0.1:8000
      </text>
    </svg>
    '''
    return base64.b64encode(placeholder_svg.encode('utf-8')).decode('utf-8'), "svg"

@app.get("/api/export/pdf")
async def export_pdf():
    from fastapi.responses import FileResponse
    import urllib.parse
    import base64
    import requests
    import shutil
    import pdfkit
    import history
    
    recent = history.get_recent_analyses(limit=1)
    if not recent:
        raise HTTPException(status_code=404, detail="No report found. Please run analysis first.")
        
    last_id = recent[0]["id"]
    data = history.get_analysis_by_id(last_id)
    if not data:
        raise HTTPException(status_code=404, detail="Report details not found.")
        
    idea = data["idea"]
    heatmap = data["heatmap"]
    fixes = data["fixes"]
    graveyard_text = data["graveyard"]
    md_content = data["raw_content"]
    scraped_urls = data.get("scraped_urls", [])

    # 1. Parse Mermaid flowchart block
    mermaid_code = ""
    mermaid_match = re.search(r'```mermaid\s*([\s\S]*?)```', md_content)
    if mermaid_match:
        mermaid_code = mermaid_match.group(1).strip()
            
    # Fetch visual flowchart
    img_tag = ""
    if mermaid_code:
        b64_data, mime_type = render_mermaid_to_image(mermaid_code)
        img_tag = f'<div style="text-align: center; margin: 15px 0;"><img src="data:image/{mime_type};base64,{b64_data}" style="max-width: 100%; height: auto;"></div>'

    # Build sources table
    sources_html = ""
    if scraped_urls:
        sources_html += """
        <div class="section-title">🌐 Sources Scraped</div>
        <div style="font-size: 9.5pt; color: #475569; margin-top: 5px;">
            <table style="width: 100%; border-collapse: collapse; margin-top: 8px;">
        """
        for i, url in enumerate(scraped_urls):
            from urllib.parse import urlparse
            domain = urlparse(url).netloc or f"Source {i+1}"
            sources_html += f"""
                <tr style="border-bottom: 1px solid #e2e8f0;">
                    <td style="padding: 8px 0; font-weight: bold; width: 45%;">{domain}</td>
                    <td style="padding: 8px 0; font-family: monospace; font-size: 8.5pt;"><a href="{url}" style="color: #dc2626; text-decoration: none;">{url}</a></td>
                </tr>
            """
        sources_html += """
            </table>
        </div>
        """

    # Compile dynamic Executive Summary
    cost_risks = [f["issue"] for f in fixes if f.get("category") == "Cost"]
    text_risks = [f["issue"] for f in fixes if f.get("category") == "Tech"]
    ux_risks = [f["issue"] for f in fixes if f.get("category") == "UX"]
    pivots = [f["issue"] for f in fixes if f.get("category") == "General"]
    
    summary_text = ""
    all_risks_text = ", ".join(ux_risks + text_risks + cost_risks)
    if all_risks_text:
        summary_text += f"The proposed idea faces several critical threats: {all_risks_text}. "
    else:
        summary_text += "The proposed idea faces low direct market friction, but you should proceed with caution. "
    
    pivots_text = " and ".join([p.replace("Pivot Strategy ", "").replace("Pivot Strategy 1: ", "").replace("Pivot Strategy 2: ", "") for p in pivots])
    if pivots_text:
        summary_text += f"To survive and capture value, consider these pivot options: {pivots_text}."

    # Severity mappings
    def get_severity(value):
        return "Very High" if value >= 3 else "High" if value == 2 else "Medium" if value == 1 else "Low"
        
    def get_color(value):
        return "#ef4444" if value >= 3 else "#f97316" if value == 2 else "#eab308" if value == 1 else "#10b981"
        
    cost_val = heatmap.get("Cost", 0)
    tech_val = heatmap.get("Tech", 0)
    ux_val = heatmap.get("UX", 0)
    
    cost_sev = get_severity(cost_val)
    tech_sev = get_severity(tech_val)
    ux_sev = get_severity(ux_val)
    
    cost_color = get_color(cost_val)
    tech_color = get_color(tech_val)
    ux_color = get_color(ux_val)
    
    # Format actionable fixes list
    fixes_html = ""
    for idx, fix in enumerate(fixes):
        title = ""
        desc = fix.get("issue", "")
        category = fix.get("category", "General")
        if category == "General":
            parts = desc.split(":")
            if len(parts) > 1:
                title = parts[0].strip()
                desc = parts[1].strip()
            else:
                title = "Pivot Strategy"
        else:
            title = f"Mitigate {category} Failure"
            
        fixes_html += f"""
        <div style="margin-bottom: 12px; padding: 10px; background-color: #f8fafc; border-left: 4px solid #64748b; border-radius: 4px;">
            <div style="font-weight: bold; font-size: 11pt; color: #0f172a;">Option {idx + 1}: {title}</div>
            <div style="font-size: 10pt; color: #475569; margin-top: 4px;">{desc}</div>
        </div>
        """
        
    # Styled HTML Template
    styled_html = f"""
    <html>
    <head>
    <style>
        @page {{
            size: letter;
            margin: 0.8in;
        }}
        body {{
            font-family: Helvetica, Arial, sans-serif;
            color: #1e293b;
            line-height: 1.5;
        }}
        .header {{
            border-bottom: 3px solid #dc2626;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }}
        .title {{
            font-size: 22pt;
            font-weight: bold;
            color: #dc2626;
        }}
        .subtitle {{
            font-size: 10pt;
            color: #64748b;
            margin-top: 5px;
            text-transform: uppercase;
        }}
        .section-title {{
            font-size: 13pt;
            font-weight: bold;
            color: #0f172a;
            margin-top: 20px;
            margin-bottom: 10px;
            border-bottom: 1px solid #e2e8f0;
            padding-bottom: 4px;
        }}
        .summary-box {{
            background-color: #ffedd5;
            border: 1px solid #fed7aa;
            padding: 12px;
            border-radius: 6px;
            font-size: 10pt;
            color: #9a3412;
            margin-bottom: 15px;
        }}
        .risk-table {{
            width: 100%;
            margin-bottom: 15px;
        }}
        .risk-card {{
            width: 31%;
            padding: 12px;
            border-radius: 6px;
            border: 1px solid #e2e8f0;
            background-color: #f8fafc;
            text-align: center;
        }}
        .graveyard {{
            background-color: #fef2f2;
            border: 1px solid #fee2e2;
            border-left: 4px solid #ef4444;
            padding: 12px;
            border-radius: 6px;
            font-family: Courier, monospace;
            font-size: 9.5pt;
            color: #991b1b;
            margin-bottom: 15px;
        }}
    </style>
    </head>
    <body>
        <div class="header">
            <div class="title">🐦🔥 PhoenixForge Autopsy Report</div>
            <div class="subtitle">Idea: {idea} &bull; Generated on: {datetime.now().strftime('%Y-%m-%d')}</div>
        </div>
        
        <div class="section-title">🔍 Executive Summary</div>
        <div class="summary-box">
            {summary_text}
        </div>
        
        <div class="section-title">📊 Risk Assessment Heatmap</div>
        <table class="risk-table">
            <tr>
                <td class="risk-card" style="border-top: 4px solid {cost_color};">
                    <div style="font-size: 9pt; color: #64748b; font-weight: bold; text-transform: uppercase;">Financial & Scaling</div>
                    <div style="font-size: 14pt; font-weight: bold; color: {cost_color}; margin-top: 6px;">{cost_sev}</div>
                    <div style="font-size: 8pt; color: #94a3b8; margin-top: 4px;">{cost_val} risk signals detected</div>
                </td>
                <td style="width: 3.5%;"></td>
                <td class="risk-card" style="border-top: 4px solid {tech_color};">
                    <div style="font-size: 9pt; color: #64748b; font-weight: bold; text-transform: uppercase;">Market & B2B</div>
                    <div style="font-size: 14pt; font-weight: bold; color: {tech_color}; margin-top: 6px;">{tech_sev}</div>
                    <div style="font-size: 8pt; color: #94a3b8; margin-top: 4px;">{tech_val} risk signals detected</div>
                </td>
                <td style="width: 3.5%;"></td>
                <td class="risk-card" style="border-top: 4px solid {ux_color};">
                    <div style="font-size: 9pt; color: #64748b; font-weight: bold; text-transform: uppercase;">UX & Retention</div>
                    <div style="font-size: 14pt; font-weight: bold; color: {ux_color}; margin-top: 6px;">{ux_sev}</div>
                    <div style="font-size: 8pt; color: #94a3b8; margin-top: 4px;">{ux_val} risk signals detected</div>
                </td>
            </tr>
        </table>
        
        <div class="section-title">💀 The Graveyard (Negative Signals)</div>
        <div style="font-size: 8.5pt; color: #64748b; margin-bottom: 6px; font-style: italic;">Scraped user complaints and autopsies:</div>
        <div class="graveyard">
            {graveyard_text.replace('\n', '<br>')}
        </div>
        
        <div class="section-title">🗺️ System Flowchart (Architecture Risks)</div>
        {img_tag}
        
        <div class="section-title">🩹 Actionable Fixes & Pivot Options</div>
        <div style="margin-top: 8px;">
            {fixes_html}
        </div>
        
        {sources_html}
        
        <div style="text-align: right; font-size: 8pt; color: #94a3b8; margin-top: 30px; border-top: 1px solid #e2e8f0; padding-top: 8px;">
            Page <pdf:pagenumber>
        </div>
    </body>
    </html>
    """
    
    pdf_path = "phoenixforge_report.pdf"
    pdf_generated = False
    
    wkhtml_path = shutil.which("wkhtmltopdf")
    if not wkhtml_path:
        common_paths = [
            r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe",
            r"C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe"
        ]
        for p in common_paths:
            if os.path.exists(p):
                wkhtml_path = p
                break
                
    if wkhtml_path:
        try:
            print(f"Generating PDF with pdfkit using wkhtmltopdf from: {wkhtml_path}")
            config = pdfkit.configuration(wkhtmltopdf=wkhtml_path)
            options = {
                'page-size': 'Letter',
                'margin-top': '0.8in',
                'margin-right': '0.8in',
                'margin-bottom': '0.8in',
                'margin-left': '0.8in',
                'encoding': "UTF-8",
                'no-outline': None,
                'enable-local-file-access': None
            }
            pdfkit.from_string(styled_html, pdf_path, configuration=config, options=options)
            pdf_generated = True
        except Exception as e:
            print(f"pdfkit failed: {e}. Falling back to xhtml2pdf...")
            
    if not pdf_generated:
        print("Generating PDF via xhtml2pdf fallback...")
        from xhtml2pdf import pisa
        with open(pdf_path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(styled_html, dest=result_file)
        if pisa_status.err:
            raise HTTPException(status_code=500, detail="Failed to generate PDF via both pdfkit and xhtml2pdf")
        
    return FileResponse(pdf_path, media_type="application/pdf", filename="phoenixforge_report.pdf")

@app.get("/api/export/word")
async def export_word():
    from docx import Document
    from docx.shared import Inches
    from fastapi.responses import FileResponse
    import history
    
    recent = history.get_recent_analyses(limit=1)
    if not recent:
        raise HTTPException(status_code=404, detail="No report found. Please run analysis first.")
        
    last_id = recent[0]["id"]
    data = history.get_analysis_by_id(last_id)
    if not data:
        raise HTTPException(status_code=404, detail="Report details not found.")
        
    content = data["raw_content"]
    scraped_urls = data.get("scraped_urls", [])
        
    # Strip pipeline metadata
    if "## Pipeline Metadata" in content:
        content = content.split("## Pipeline Metadata")[0]
    
    doc = Document()
    doc.add_heading("PhoenixForge Report", level=1)
    
    for line in content.split("\n"):
        stripped_line = line.strip()
        if not stripped_line:
            continue
        if stripped_line.startswith("# "):
            doc.add_heading(stripped_line[2:], level=1)
        elif stripped_line.startswith("## "):
            doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith("### "):
            doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith("```") or stripped_line.startswith("`"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(stripped_line)
            run.font.name = 'Courier New'
        else:
            doc.add_paragraph(stripped_line)
            
    # Add sources section to word document
    if scraped_urls:
        doc.add_heading("Sources Scraped", level=2)
        for url in scraped_urls:
            doc.add_paragraph(url)
            
    export_path = "phoenixforge_report.docx"
    doc.save(export_path)
    return FileResponse(export_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="phoenixforge_report.docx")

@app.get("/api/history")
async def get_history():
    try:
        import history
        recent = history.get_recent_analyses(limit=10)
        return {"status": "success", "history": recent}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/history/{id}")
async def get_history_by_id(id: int):
    try:
        import history
        data = history.get_analysis_by_id(id)
        if not data:
            raise HTTPException(status_code=404, detail="History entry not found")
        return {
            "status": "success",
            "heatmap": data["heatmap"],
            "fixes": data["fixes"],
            "graveyard": data["graveyard"],
            "pipeline": extract_pipeline_metadata(data["raw_content"]),
            "raw_content": data["raw_content"],
            "raw_combined_markdown": data.get("raw_combined_markdown", ""),
            "scraped_urls": data.get("scraped_urls", []),
            "source_count": data.get("source_count", 0)
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Mount static files to serve the frontend (must be defined after endpoints)
app.mount("/", StaticFiles(directory="static", html=True), name="static")
