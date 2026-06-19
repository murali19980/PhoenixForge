"""
documents.py - Structured document generation for PhoenixForge
Generates Project Charter, Project Management Plan, and Complete Report
in both PDF and Word formats.
"""

import io
import re
import logging
from datetime import datetime
from typing import Dict, Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# Set up logging
logger = logging.getLogger("phoenixforge.documents")

# --- Core Data Extraction ---

def extract_analysis_data(raw_content: str, heatmap: Dict, fixes: list, graveyard: str, idea: str = "Unknown Project") -> Dict[str, Any]:
    """Extract structured data from analysis results."""
    # Prioritize idea parameter, fallback to content parsing
    project_idea = idea
    if not project_idea or project_idea == "Unknown Project":
        first_line = raw_content.split('\n')[0] if raw_content else ""
        if 'Report:' in first_line:
            project_idea = first_line.split('Report:')[-1].strip()
        elif 'Report for' in first_line:
            project_idea = first_line.split('Report for')[-1].strip()

    # Count risks by category
    risk_counts = {"UX": 0, "Tech": 0, "Cost": 0}
    for fix in fixes:
        cat = fix.get('category', '')
        if cat == 'UX':
            risk_counts['UX'] += 1
        elif cat == 'Tech' or cat == 'Market':
            risk_counts['Tech'] += 1
        elif cat == 'Cost':
            risk_counts['Cost'] += 1
    
    # Extract pivot strategies
    pivots = [f for f in fixes if f.get('category') == 'General']
    
    # Extract metrics (percentages, dollar amounts, etc.)
    metrics = re.findall(r'\b\d+(?:\.\d+)?\s*(?:%|\$|usd|users|requests|customers|clients)\b', raw_content, re.IGNORECASE)
    
    # Extract key risks (first 3 from each category)
    key_risks = {
        "UX": [f.get('issue', '') for f in fixes if f.get('category') == 'UX'][:3],
        "Tech": [f.get('issue', '') for f in fixes if f.get('category') in ('Tech', 'Market')][:3],
        "Cost": [f.get('issue', '') for f in fixes if f.get('category') == 'Cost'][:3],
    }
    
    return {
        "idea": project_idea,
        "date": datetime.now().strftime("%B %d, %Y"),
        "risk_counts": risk_counts,
        "key_risks": key_risks,
        "pivots": pivots,
        "metrics": metrics[:5],
        "graveyard": graveyard if graveyard else "No negative signals found.",
        "total_risks": len(fixes),
    }


# --- Markdown to Word Helpers ---

def add_markdown_to_doc(doc: Document, text: str):
    """Convert simple markdown to Word document elements."""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:])
            p.style = doc.styles['Heading 1']
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.style = doc.styles['Heading 2']
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style = doc.styles['Heading 3']
        # Bullet points
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('•'):
            clean_line = line[1:].strip()
            p = doc.add_paragraph(clean_line, style='List Bullet')
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            if p.runs:
                p.runs[0].italic = True
        else:
            doc.add_paragraph(line)


# --- Project Charter Generator ---

def generate_charter(data: Dict[str, Any]) -> Document:
    """Generate a Project Charter document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Charter', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(f'Project Name: {data["idea"]}', level=1)
    doc.add_paragraph(f'Date of Charter: {data["date"]}')
    doc.add_paragraph()
    
    # 1. Purpose / Business Case
    doc.add_heading('1. Purpose & Business Case', level=2)
    doc.add_paragraph(
        f'This charter formalizes the intent to build and validate "{data["idea"]}". '
        f'Based on our market research and competitive risk autopsies (which identified {data["total_risks"]} total failure points), '
        f'the project serves to solve key failures that commonly crash similar platforms. Key risks to mitigate from day one include:'
    )
    for risk in data['key_risks'].get('UX', [])[:2]:
        doc.add_paragraph(f'{risk}', style='List Bullet')
    for risk in data['key_risks'].get('Tech', [])[:2]:
        doc.add_paragraph(f'{risk}', style='List Bullet')
    doc.add_paragraph()
    
    # 2. Objectives & Success Criteria
    doc.add_heading('2. Objectives & Success Criteria', level=2)
    doc.add_paragraph(
        'The primary objectives and success milestones of the project are defined as follows:'
    )
    success_criteria = [
        'Mitigate identified user onboarding and retention failure loops.',
        'Establish a robust and budget-conscious technical stack (under $50/mo initial).',
        'Verify quantitative metrics and market signals before full-scale dev budget release.',
        'Validate product-market fit using a minimal viable prototype within 8 weeks.'
    ]
    for criterion in success_criteria:
        doc.add_paragraph(criterion, style='List Bullet')
    doc.add_paragraph()
    
    # 3. Project Scope
    doc.add_heading('3. Project Scope', level=2)
    doc.add_paragraph('In-Scope Items:')
    doc.add_paragraph('• Core feature set engineering and prototyping', style='List Bullet')
    doc.add_paragraph('• Basic onboarding workflow and feedback loops', style='List Bullet')
    doc.add_paragraph('• Primary failure node mitigations (as mapped by PhoenixForge)', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Out-of-Scope Items:')
    doc.add_paragraph('• Multi-region database replication and scaling', style='List Bullet')
    doc.add_paragraph('• High-budget paywall marketing and customer acquisition campaigns', style='List Bullet')
    doc.add_paragraph()
    
    # 4. Stakeholders
    doc.add_heading('4. Key Stakeholders', level=2)
    doc.add_paragraph('• Executive Sponsor: [To be assigned]')
    doc.add_paragraph('• Project Manager / Owner: [To be assigned]')
    doc.add_paragraph('• Lead Engineer: [To be assigned]')
    doc.add_paragraph('• Target Audience: Early adopters and target user segments')
    doc.add_paragraph()
    
    # 5. High-Level Risks
    doc.add_heading('5. High-Level Risks & Quantified Factors', level=2)
    doc.add_paragraph('The PhoenixForge scan has flagged the following high-level risk categories:')
    
    risk_descriptions = {
        'UX': 'User experience, onboarding, and churn risks',
        'Tech': 'Technical execution, API limits, and scaling risks',
        'Cost': 'Resource burn-rate and operational margins'
    }
    for category, count in data['risk_counts'].items():
        if count > 0:
            doc.add_paragraph(
                f'{risk_descriptions.get(category, category)}: {count} risk signals detected',
                style='List Bullet'
            )
    doc.add_paragraph()
    
    # 6. Timeline Estimates
    doc.add_heading('6. High-Level Timeline', level=2)
    doc.add_paragraph('• Discovery & Risk Planning: Weeks 1 - 2')
    doc.add_paragraph('• Prototype Development: Weeks 3 - 6')
    doc.add_paragraph('• Quality Assurance & Launch Prep: Weeks 7 - 8')
    doc.add_paragraph()
    
    # 7. Approvals
    doc.add_heading('7. Authorizations & Sign-off', level=2)
    doc.add_paragraph('Signatures below approve this Project Charter and authorize the allocation of validation resources.')
    doc.add_paragraph()
    doc.add_paragraph('Approved by: ________________________      Date: ________________________')
    
    return doc


# --- Project Management Plan Generator ---

def generate_plan(data: Dict[str, Any]) -> Document:
    """Generate a Project Management Plan document."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Management Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(f'Project: {data["idea"]}', level=1)
    doc.add_paragraph(f'Date: {data["date"]}')
    doc.add_paragraph()
    
    # 1. Scope Management
    doc.add_heading('1. Scope Management', level=2)
    doc.add_paragraph('The team will build and deploy a solution that matches the scoped requirements while remaining strictly aligned with our risk mitigation plan:')
    doc.add_paragraph(
        f'Project Scope Statement: Deliver a minimal viable version of "{data["idea"]}" '
        f'that implements core safeguards against identified failure loops.'
    )
    doc.add_paragraph()
    doc.add_paragraph('Core Work Breakdown Structure (WBS):')
    doc.add_paragraph('1. Planning & Risk Assessment', style='List Bullet')
    doc.add_paragraph('2. UX Design & Wireframing', style='List Bullet')
    doc.add_paragraph('3. Coding & API Integration', style='List Bullet')
    doc.add_paragraph('4. Quality Testing', style='List Bullet')
    doc.add_paragraph('5. Deployment & Release', style='List Bullet')
    doc.add_paragraph()
    
    # 2. Schedule Management
    doc.add_heading('2. Schedule & Milestones', level=2)
    doc.add_paragraph('Key schedule milestones for execution tracking:')
    milestones = [
        ('M1', 'Discovery & Planning Complete', 'Week 2'),
        ('M2', 'Core Features Coded', 'Week 6'),
        ('M3', 'Testing & Bug Fixes Done', 'Week 8'),
        ('M4', 'Initial Release & Feedback Loop', 'Week 9'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Milestone Description'
    hdr_cells[2].text = 'Target Week'
    for mid, name, date in milestones:
        row_cells = table.add_row().cells
        row_cells[0].text = mid
        row_cells[1].text = name
        row_cells[2].text = date
    doc.add_paragraph()
    
    # 3. Resource Management
    doc.add_heading('3. Resource & Team Roles', level=2)
    doc.add_paragraph('• Project Manager: Overall coordinator')
    doc.add_paragraph('• Lead Developer: Technical execution and API design')
    doc.add_paragraph('• QA / Testing: Validates features against UX checklists')
    doc.add_paragraph()
    
    # 4. Risk Register & Mitigation Plans
    doc.add_heading('4. Risk Register & Mitigation', level=2)
    doc.add_paragraph('Detailed Failure Risks and Mitigations (Identified via scans):')
    for category, count in data['risk_counts'].items():
        if count > 0:
            doc.add_paragraph(f'{category} Failures ({count} signals):', style='List Bullet')
            for risk in data['key_risks'].get(category, [])[:3]:
                doc.add_paragraph(f'  - {risk}', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('Mitigation Policy: Every risk listed in the register must have an automated test or a design standard in the codebase before code approval.')
    doc.add_paragraph()
    
    # 5. Quality Management
    doc.add_heading('5. Quality Management', level=2)
    doc.add_paragraph('• All code changes require peer review.')
    doc.add_paragraph('• Core onboarding flows must undergo manual UX testing.')
    doc.add_paragraph('• API rate limits must be monitored.')
    doc.add_paragraph()
    
    # 6. Communications Plan
    doc.add_heading('6. Communications Plan', level=2)
    doc.add_paragraph('• Daily: Core team standup (15 mins)')
    doc.add_paragraph('• Weekly: Risk and sprint status review')
    doc.add_paragraph('• Post-Launch: Bi-weekly cohort retention analysis')
    doc.add_paragraph()
    
    # 7. Pivot Options
    if data['pivots']:
        doc.add_heading('7. Strategic Pivot Guidelines', level=2)
        doc.add_paragraph('If the core idea hits high retention churn, consider pivoting to these alternatives:')
        for pivot in data['pivots'][:3]:
            issue_text = pivot.get("issue", "")
            if ":" in issue_text:
                issue_text = issue_text.split(":", 1)[1].strip()
            doc.add_paragraph(issue_text, style='List Bullet')
        doc.add_paragraph()
    
    return doc


# --- Combined Report Generator ---

def generate_complete_report(data: Dict[str, Any]) -> Document:
    """Generate a Complete Report (Charter + Plan + Analysis)."""
    doc = Document()
    
    # Title Page
    title = doc.add_heading('PhoenixForge Complete Risk & Execution Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f'Project: {data["idea"]}', level=1)
    doc.add_paragraph(f'Generated on: {data["date"]}')
    doc.add_page_break()
    
    # Part 1: Project Charter
    doc.add_heading('PART 1: PROJECT CHARTER', level=1)
    charter_doc = generate_charter(data)
    for element in charter_doc.element.body:
        doc.element.body.append(element)
    doc.add_page_break()
    
    # Part 2: Project Management Plan
    doc.add_heading('PART 2: PROJECT MANAGEMENT PLAN', level=1)
    plan_doc = generate_plan(data)
    for element in plan_doc.element.body:
        doc.element.body.append(element)
    doc.add_page_break()
    
    # Part 3: Graveyard (Raw Scraped Analysis)
    doc.add_heading('PART 3: RISK RESEARCH & GRAVEYARD', level=1)
    doc.add_paragraph("The following raw failure logs, competitive autopsies, and source signals were scraped during discovery:")
    doc.add_paragraph()
    add_markdown_to_doc(doc, data.get('graveyard', 'No analysis data available.'))
    
    return doc


# --- PDF Generation (using reportlab) ---

def docx_to_pdf(doc: Document) -> bytes:
    """Convert a Word document structure to a styled PDF using ReportLab layout objects."""
    buffer = io.BytesIO()
    
    # Page setup
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    # Professional layout styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor('#0f172a'),
        alignment=1, # Center
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'DocH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#0f172a'),
        spaceBefore=15,
        spaceAfter=10,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#334155'),
        spaceBefore=12,
        spaceAfter=8,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#334155'),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )
    
    story = []
    
    # Iterate elements in document body to preserve exact paragraph & table order
    for element in doc.element.body:
        if element.tag.endswith('p'):
            from docx.text.paragraph import Paragraph as DocxParagraph
            paragraph = DocxParagraph(element, doc)
            text = paragraph.text.strip()
            if not text:
                continue
                
            style_name = paragraph.style.name.lower()
            
            # Map Word style names to reportlab styles
            if 'title' in style_name or 'heading 0' in style_name:
                story.append(Paragraph(text, title_style))
                story.append(Spacer(1, 10))
            elif 'heading 1' in style_name:
                story.append(Paragraph(text, h1_style))
            elif 'heading 2' in style_name:
                story.append(Paragraph(text, h2_style))
            elif 'heading 3' in style_name:
                story.append(Paragraph(text, h2_style))
            elif 'bullet' in style_name or text.startswith('•') or text.startswith('-'):
                bullet_text = text
                if bullet_text.startswith('•') or bullet_text.startswith('-'):
                    bullet_text = bullet_text[1:].strip()
                story.append(Paragraph(f"&bull; {bullet_text}", bullet_style))
            else:
                story.append(Paragraph(text, body_style))
                
        elif element.tag.endswith('tbl'):
            from docx.table import Table as DocxTable
            table = DocxTable(element, doc)
            
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(Paragraph(cell_text, body_style))
                table_data.append(row_data)
                
            if table_data:
                # Column widths equally divided
                col_count = len(table_data[0])
                col_width = (8.5 * inch - 2 * 0.75 * inch) / col_count
                
                rl_table = Table(table_data, colWidths=[col_width]*col_count)
                rl_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f1f5f9')),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#0f172a')),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 10),
                    ('BOTTOMPADDING', (0,0), (-1,0), 8),
                    ('TOPPADDING', (0,0), (-1,0), 8),
                    ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#ffffff')),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ('BOTTOMPADDING', (0,1), (-1,-1), 6),
                    ('TOPPADDING', (0,1), (-1,-1), 6),
                ]))
                story.append(Spacer(1, 10))
                story.append(rl_table)
                story.append(Spacer(1, 10))
                
        elif element.tag.endswith('br') or 'pagebreak' in str(element):
            story.append(PageBreak())
            
    pdf.build(story)
    buffer.seek(0)
    return buffer.read()


# --- Main Export Functions ---

def generate_document(data: Dict[str, Any], doc_type: str, format: str) -> bytes:
    """Main dispatcher generating DOCX and PDF byte streams."""
    if doc_type == 'charter':
        doc = generate_charter(data)
    elif doc_type == 'plan':
        doc = generate_plan(data)
    elif doc_type == 'complete':
        doc = generate_complete_report(data)
    else:
        raise ValueError(f"Unknown document type: {doc_type}")
    
    if format == 'pdf':
        return docx_to_pdf(doc)
    elif format == 'word':
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    else:
        raise ValueError(f"Unknown format: {format}")
